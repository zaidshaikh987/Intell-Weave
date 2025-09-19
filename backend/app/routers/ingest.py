"""
app/routers/ingest.py
- Real extraction endpoints for URL and file uploads.
- Uses readability + BeautifulSoup and feeds NLP pipeline before persisting.
"""
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import text
from sqlalchemy.dialects.postgresql import insert as pg_insert
from ..core.db import get_db
from ..models.tables import articles, article_nlp
from ..services.nlp import nlp
from bs4 import BeautifulSoup
from readability import Document
from urllib.parse import urlparse, parse_qsl, urlencode, urlunparse
from datetime import datetime, timezone
import requests
import io
from langdetect import detect as lang_detect

router = APIRouter()

class UrlIn(BaseModel):
    url: str


def _strip_utm(u: str) -> str:
    pr = urlparse(u)
    q = [(k, v) for k, v in parse_qsl(pr.query) if not k.lower().startswith('utm_')]
    return urlunparse((pr.scheme, pr.netloc, pr.path, pr.params, urlencode(q), pr.fragment))

@router.post("/url")
def ingest_url(body: UrlIn, db: Session = Depends(get_db)):
    # fetch
    try:
        resp = requests.get(body.url, timeout=15)
        resp.raise_for_status()
        html = resp.text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"failed to fetch url: {e}")

    # parse readability
    try:
        doc = Document(html)
        title = (doc.short_title() or '').strip()
        content_html = doc.summary(html_partial=True)
        soup_read = BeautifulSoup(content_html, 'lxml')
        paragraphs = [p.get_text(" ", strip=True) for p in soup_read.find_all('p')]
        body_text = "\n".join(paragraphs)[:20000]
        body_html = str(soup_read)[:50000]
        subtitle_tag = soup_read.find(['h2','h3'])
        subtitle = subtitle_tag.get_text(" ", strip=True)[:300] if subtitle_tag else None
    except Exception:
        soup = BeautifulSoup(html, 'lxml')
        title = (soup.title.string or '').strip() if soup.title else body.url
        paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all('p')]
        body_text = "\n".join(paragraphs)[:20000]
        body_html = None
        subtitle_tag = soup.find(['h2','h3'])
        subtitle = subtitle_tag.get_text(" ", strip=True)[:300] if subtitle_tag else None

    # meta
    soup_meta = BeautifulSoup(html, 'lxml')
    canon = soup_meta.find('link', rel=lambda v: v and 'canonical' in v)
    canonical_url = _strip_utm(canon['href'].strip()) if canon and canon.has_attr('href') else _strip_utm(body.url)
    pub = None
    for sel in [
        {'name': 'meta', 'key': 'property', 'val': 'article:published_time'},
        {'name': 'meta', 'key': 'name', 'val': 'pubdate'},
        {'name': 'meta', 'key': 'name', 'val': 'date'},
        {'name': 'meta', 'key': 'itemprop', 'val': 'datePublished'},
    ]:
        tag = soup_meta.find(sel['name'], attrs={sel['key']: sel['val']})
        if tag and tag.get('content'):
            try:
                pub = datetime.fromisoformat(tag['content'].replace('Z', '+00:00'))
                break
            except Exception:
                continue
    meta_kw = soup_meta.find('meta', attrs={'name':'keywords'})
    tags = [t.strip() for t in (meta_kw['content'].split(',') if meta_kw and meta_kw.get('content') else []) if t.strip()] or None

    # author detection
    author = None
    for sel in [
        {'name': 'meta', 'key': 'name', 'val': 'author'},
        {'name': 'meta', 'key': 'property', 'val': 'article:author'},
        {'name': 'meta', 'key': 'name', 'val': 'byl'},
    ]:
        tag = soup_meta.find(sel['name'], attrs={sel['key']: sel['val']})
        if tag and tag.get('content'):
            author = tag['content'].strip()
            break

    # language detection and reading time
    try:
        language = lang_detect(body_text) if body_text and len(body_text)>40 else None
    except Exception:
        language = None
    words = len(body_text.split()) if body_text else 0
    reading_time = max(1, words // 200) if words else None

    # best image (og:image or first content image)
    image_url = None
    ogimg = soup_meta.find('meta', attrs={'property':'og:image'}) or soup_meta.find('meta', attrs={'name':'og:image'})
    if ogimg and ogimg.get('content'):
        image_url = ogimg['content'].strip()
    if not image_url and 'soup_read' in locals():
        img = soup_read.find('img')
        if img and img.get('src'):
            image_url = img['src']

    # persist
    aid = f"art_{abs(hash(canonical_url))}"
    analysis = nlp.analyze(body_text)
    try:
        db.execute(
            pg_insert(articles)
            .values(
                id=aid,
                title=title or canonical_url,
                subtitle=subtitle,
                author=author,
                source_url=body.url,
                canonical_url=canonical_url,
                language=language,
                body_text=body_text,
                body_html=body_html,
                tags=tags,
                reading_time=reading_time,
                published_at=pub,
            )
            .on_conflict_do_nothing(index_elements=[articles.c.id])
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB insert failed (articles): {e}")
    # if conflict nothing inserted; still upsert NLP
    try:
        db.execute(
            pg_insert(article_nlp)
            .values(
                article_id=aid,
                summary=analysis.get('summary'),
                sentiment=analysis.get('sentiment'),
                key_entities=analysis.get('entities'),
                topics=analysis.get('keyphrases'),
                credibility_score=70.0,
            )
            .on_conflict_do_nothing(index_elements=[article_nlp.c.article_id])
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB insert failed (article_nlp): {e}")
    try:
        emb = analysis.get('embedding', [])
        vec_literal = f"[{','.join(str(round(float(x),6)) for x in emb)}]"
        db.execute(text("""
            UPDATE article_nlp SET embedding = :vec::vector
            WHERE article_id = :aid
        """), {"vec": vec_literal, "aid": aid})
    except Exception:
        pass
    db.commit()
    return {"id": aid, "title": title, "content": body_text, "source_url": body.url, "summary": analysis.get('summary'), "sentiment": analysis.get('sentiment'), "topics": analysis.get('keyphrases'), "reading_time": reading_time, "image_url": image_url, "language": language, "author": author}


@router.post("/file")
def ingest_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Read file content
    raw = file.file.read()
    text_content = ""
    name = (file.filename or "").lower()
    try:
        if name.endswith('.pdf'):
            # lightweight pdf text extraction
            try:
                from pdfminer.high_level import extract_text
                text_content = extract_text(io.BytesIO(raw))
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"failed to parse pdf: {e}")
        elif name.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(io.BytesIO(raw))
                text_content = "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"failed to parse docx: {e}")
        elif name.endswith('.html') or name.endswith('.htm'):
            html = raw.decode('utf-8', errors='ignore')
            doc = Document(html)
            content_html = doc.summary(html_partial=True)
            soup_read = BeautifulSoup(content_html, 'lxml')
            text_content = "\n".join([p.get_text(" ", strip=True) for p in soup_read.find_all('p')])
        else:
            # assume plain text
            text_content = raw.decode('utf-8', errors='ignore')
    finally:
        file.file.close()

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="no text content extracted")

    # Persist via create flow
    analysis = nlp.analyze(text_content)
    reading_time = max(1, len(text_content)//1000)
    aid = f"art_{int(datetime.now(tz=timezone.utc).timestamp()*1000)}"
    try:
        db.execute(
            pg_insert(articles)
            .values(id=aid, title=(file.filename or 'Uploaded Article'), body_text=text_content, reading_time=reading_time)
            .on_conflict_do_nothing(index_elements=[articles.c.id])
        )
        db.execute(
            pg_insert(article_nlp)
            .values(article_id=aid, summary=analysis.get('summary'), sentiment=analysis.get('sentiment'), key_entities=analysis.get('entities'), topics=analysis.get('keyphrases'), credibility_score=70.0)
            .on_conflict_do_nothing(index_elements=[article_nlp.c.article_id])
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB insert failed: {e}")
    try:
        emb = analysis.get('embedding', [])
        vec_literal = f"[{','.join(str(round(float(x),6)) for x in emb)}]"
        db.execute(text("""
            UPDATE article_nlp SET embedding = :vec::vector
            WHERE article_id = :aid
        """), {"vec": vec_literal, "aid": aid})
    except Exception:
        pass
    db.commit()
    return {"id": aid, "title": file.filename or 'Uploaded Article', "content": text_content, "summary": analysis.get('summary'), "topics": analysis.get('keyphrases'), "sentiment": analysis.get('sentiment'), "reading_time": reading_time}
